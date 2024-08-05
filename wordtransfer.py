import os
import json
import re
import time
from docx import Document

def clean_text(text):
    """
    Remove non-printable characters from the text.
    """
    # Replace control characters with a space
    return re.sub(r'[\x00-\x1F\x7F-\x9F]', ' ', text)

# Start timing
start_time = time.time()

# Print starting message
print("Process started...")

# Define the file path to the JSON file in the Documents folder
home_directory = os.path.expanduser('~')
file_path = os.path.join(home_directory, 'Documents', 'data.json')

# Step 1: Read and parse the JSON data
with open(file_path, 'r') as file:
    data = [json.loads(line) for line in file]

# Step 2: Filter out entries with "error": "question_not_found"
filtered_data = [entry for entry in data if not (entry.get('error') == 'question_not_found')]

# Step 3: Organize the remaining data
# Create a list of dictionaries to hold the organized data
organized_data = []
for entry in filtered_data:
    question = entry.get('question', {})
    answers = entry.get('answers', [])
    
    question_id = question.get('id')
    created_at = question.get('created_at')
    updated_at = question.get('updated_at')
    title = clean_text(question.get('title', ''))
    detail = clean_text(question.get('detail', ''))
    
    for answer in answers:
        answer_detail = clean_text(answer.get('detail', ''))
        answer_poster = clean_text(answer.get('poster', {}).get('name', ''))
        
        organized_data.append({
            'Question ID': question_id,
            'Created At': created_at,
            'Updated At': updated_at,
            'Title': title,
            'Detail': detail,
            'Answer': answer_detail,
            'Answered By': answer_poster
        })

# Step 4: Export the organized data to a Word document
doc = Document()
doc.add_heading('Questions and Answers', 0)

for item in organized_data:
    doc.add_heading(f"Question ID: {item['Question ID']}", level=1)
    doc.add_paragraph(f"Created At: {item['Created At']}")
    doc.add_paragraph(f"Updated At: {item['Updated At']}")
    doc.add_paragraph(f"Title: {item['Title']}")
    doc.add_paragraph(f"Detail: {item['Detail']}")
    doc.add_heading('Answers:', level=2)
    doc.add_paragraph(f"Answer: {item['Answer']}")
    doc.add_paragraph(f"Answered By: {item['Answered By']}")
    doc.add_page_break()

# Save the output document in the Documents folder
output_path = os.path.join(home_directory, 'Documents', 'organized_questions_and_answers.docx')
doc.save(output_path)

# End timing
end_time = time.time()

# Print completion message and duration
print(f"Document saved successfully at {output_path}")
print(f"Time taken: {end_time - start_time} seconds")